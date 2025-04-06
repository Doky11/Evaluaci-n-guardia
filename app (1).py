
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from io import BytesIO

st.set_page_config(page_title="Informe Guardia CAES", layout="centered")
st.title("Informe Personal de Guardia - CAES")

# ENCABEZADO
st.header("Encabezado del Informe")
informante = st.text_input("Informante")
fecha = st.date_input("Fecha", value=date.today())
alumno = st.text_input("Alumno evaluado")
puesto = st.text_input("Puesto durante la guardia")

# FUNCIONES
def calificacion_letra(nota):
    if nota >= 9:
        return "A"
    elif nota >= 7:
        return "B"
    elif nota >= 5:
        return "C"
    else:
        return "D"

# CATEGORÍAS Y PREGUNTAS (no cambiadas aquí, pero se conservan internamente)
categorias = {
    "POLICÍA": ["Pregunta"]*12,
    "DISCIPLINA": ["Pregunta"]*6,
    "INTERES": ["Pregunta"]*6,
    "RESPONSABILIDAD": ["Pregunta"]*5,
    "INICIATIVA": ["Pregunta"]*5,
    "CONFIANZA EN SI MISMO": ["Pregunta"]*5,
    "ACTITUD CON LOS SUBORDINADOS": ["Pregunta"]*6,
    "ACTITUD CON EL MANDO": ["Pregunta"]*5,
    "COMPETENCIA / EFICACIA": ["Pregunta"]*5,
    "TRATO": ["Pregunta"]*5,
    "RESISTENCIA A LA FATIGA": ["Pregunta"]*4
}

# EVALUACIÓN
st.header("Evaluación por Categorías")
notas_categoria = {}
total_puntos = 0
total_preguntas = 0

for categoria, preguntas in categorias.items():
    st.subheader(categoria)
    respuestas = [st.checkbox(p, key=f"{categoria}_{i}") for i, p in enumerate(preguntas)]
    puntos = sum(respuestas)
    nota = round((puntos / len(preguntas)) * 10, 2)
    letra = calificacion_letra(nota)
    notas_categoria[categoria] = letra
    total_puntos += puntos
    total_preguntas += len(preguntas)
    st.write(f"Nota: {nota} — Calificación: {letra}")

# NOTA MEDIA
st.header("Resumen Final")
nota_media = round((total_puntos / total_preguntas) * 10, 2)
letra_media = calificacion_letra(nota_media)
st.write(f"📊 Nota media: {nota_media}")
st.write(f"🔠 Calificación final: {letra_media}")
observaciones = st.text_area("Observaciones generales / Justificación")

# GENERADOR DE INFORME
def generar_word():
    doc = Document("INFORME EN BLANCO.docx")
    for table in doc.tables:
        for row in table.rows:
            celdas = row.cells
            if not celdas or len(celdas) < 5:
                continue
            texto = celdas[0].text.strip()
            if texto == "INFORMANTE":
                celdas[1].text = informante
            elif texto == "FECHA":
                celdas[1].text = fecha.strftime("%d.%m.%Y")
            elif texto == "ALUMNO":
                celdas[1].text = alumno
            elif texto == "PUESTO":
                celdas[1].text = puesto
            elif texto in notas_categoria:
                letra = notas_categoria[texto]
                for i in range(1, 5):
                    celdas[i].text = ""  # Limpia casillas de nota
                indice_columna = {"A": 1, "B": 2, "C": 3, "D": 4}.get(letra)
                if indice_columna is not None and indice_columna < len(celdas):
                    parrafo = celdas[indice_columna].paragraphs[0]
                    parrafo.clear()
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    parrafo.add_run("X")
            elif texto.startswith("Nota media"):
                celdas[0].text = f"Nota media: {nota_media}"
            elif texto == "":
                celdas[1].text = observaciones

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if st.button("📄 Generar Informe Word"):
    archivo = generar_word()
    st.download_button(
        label="📥 Descargar Informe",
        data=archivo,
        file_name=f"Informe_{alumno.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
