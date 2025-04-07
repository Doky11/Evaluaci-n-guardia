import streamlit as st
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from io import BytesIO

st.set_page_config(page_title="Informe Guardia CAES", layout="centered")
st.title("Informe Personal de Guardia - CAES")

# --- ENCABEZADO ---
st.header("Encabezado del Informe")
informante = st.text_input("Informante")
fecha = st.date_input("Fecha", value=date.today())
alumno = st.text_input("Alumno evaluado")
puesto = st.text_input("Puesto durante la guardia")

# --- CHECKLIST DE EVALUACIÓN ---
st.header("Evaluación por Categorías")

categorias = {
    "POLICÍA": [
        "¿Cuida su imagen y presencia durante toda la guardia?",
        "¿Lleva el cabello con color natural y corte uniforme o recogido?",
        "¿Tiene las uñas cuidadas y aseadas, sin pintar?",
        "¿Está correctamente afeitado al entrar y salir de la guardia?",
        "¿Tiene los tatuajes tapados conforme a la normativa?",
        "¿Utiliza maquillaje natural, si corresponde?",
        "¿Lleva el uniforme limpio y completo durante la guardia?",
        "¿El calzado está limpio y en condiciones reglamentarias?",
        "¿Evita gestos o actitudes informales durante el servicio?",
        "¿Adopta postura corporal profesional?",
        "¿Lleva los distintivos reglamentarios bien colocados?",
        "¿Inspira respeto y seriedad con su presencia?"
    ],
    "DISCIPLINA": [
        "¿Se exige a sí mismo en el cumplimiento de sus obligaciones?",
        "¿Cumple correctamente las órdenes recibidas?",
        "¿Mantiene comportamientos disciplinados en todo momento?",
        "¿Corrige a otros con tono adecuado?",
        "¿Acude puntual a formaciones y relevos?",
        "¿Acepta cambios sin protestar?"
    ],
    "INTERES": [
        "¿Muestra entrega y disponibilidad?",
        "¿Conoce los medios y materiales asignados?",
        "¿Demuestra interés por mejorar profesionalmente?",
        "¿Pregunta ante dudas?",
        "¿Evita distracciones o uso de móvil?",
        "¿Se mantiene atento y activo?"
    ],
    "RESPONSABILIDAD": [
        "¿Asume las órdenes dadas por él ante el mando?",
        "¿Se responsabiliza de sus decisiones?",
        "¿Informa si algo no sale correctamente?",
        "¿Se anticipa a necesidades del servicio?",
        "¿Finaliza sus cometidos sin necesidad de supervisión?"
    ],
    "INICIATIVA": [
        "¿Se ofrece voluntariamente para tareas adicionales?",
        "¿Muestra motivación en sus cometidos?",
        "¿Actúa con acierto y prudencia?",
        "¿Propone ideas de mejora?",
        "¿Reacciona con iniciativa ante imprevistos?"
    ],
    "CONFIANZA EN SI MISMO": [
        "¿Toma decisiones con seguridad?",
        "¿Expresa sus ideas con claridad y respeto?",
        "¿Resuelve problemas con actitud positiva?",
        "¿Se mantiene sereno ante la presión?",
        "¿Asume responsabilidades sin temor?"
    ],
    "ACTITUD CON LOS SUBORDINADOS": [
        "¿Trata con respeto a sus subordinados?",
        "¿Transmite claramente los cometidos?",
        "¿Es ejemplo en actitud y presencia?",
        "¿Se preocupa por su bienestar y seguridad?",
        "¿Fomenta la motivación y el buen ambiente?",
        "¿Evita favoritismos?"
    ],
    "ACTITUD CON EL MANDO": [
        "¿Muestra lealtad y cooperación?",
        "¿Trata con corrección al mando?",
        "¿Expone opiniones con educación?",
        "¿Cumple lo ordenado sin excusas?",
        "¿Consulta con claridad cuando es necesario?"
    ],
    "COMPETENCIA / EFICACIA": [
        "¿Cumple objetivos en tiempo y forma?",
        "¿Realiza el trabajo con calidad?",
        "¿Conoce perfectamente sus cometidos?",
        "¿Resuelve con eficacia los imprevistos?",
        "¿Sabe priorizar tareas?"
    ],
    "TRATO": [
        "¿Se muestra afable con todos los compañeros?",
        "¿Actúa con calma en situaciones tensas?",
        "¿Evita conflictos personales?",
        "¿Fomenta un buen ambiente de trabajo?",
        "¿Escucha activamente a los demás?"
    ],
    "RESISTENCIA A LA FATIGA": [
        "¿Mantiene su rendimiento toda la guardia?",
        "¿No muestra signos de fatiga evidentes?",
        "¿Se mantiene alerta hasta el final del turno?",
        "¿Actúa profesionalmente aunque esté cansado?"
    ]
}

notas_categoria = {}
total_puntos = 0
total_preguntas = 0

for categoria, preguntas in categorias.items():
    st.subheader(categoria)
    respuestas = [st.checkbox(preg, key=f"{categoria}_{i}") for i, preg in enumerate(preguntas)]
    puntos = sum(respuestas)
    nota = round((puntos / len(preguntas)) * 10, 2)
    total_puntos += puntos
    total_preguntas += len(preguntas)

    if nota >= 9:
        letra = "A"
    elif nota >= 7:
        letra = "B"
    elif nota >= 5:
        letra = "C"
    else:
        letra = "D"

    notas_categoria[categoria] = letra
    st.write(f"Nota: {nota} — Calificación: {letra}")

# --- NOTA MEDIA Y OBSERVACIONES ---
st.header("Resumen Final")
nota_media = round((total_puntos / total_preguntas) * 10, 2)
if nota_media >= 9:
    letra_final = "A"
elif nota_media >= 7:
    letra_final = "B"
elif nota_media >= 5:
    letra_final = "C"
else:
    letra_final = "D"

st.write(f"📊 Nota media: {nota_media}")
st.write(f"🔠 Calificación final: {letra_final}")
observaciones = st.text_area("Observaciones generales / Justificación")

# --- GENERADOR DE INFORME ---
def generar_word():
    plantilla = "INFORME EN BLANCO.docx"
    doc = Document(plantilla)

    for table in doc.tables:
        for row in table.rows:
            celdas = row.cells
            if "INFORMANTE" in celdas[0].text:
                celdas[1].text = informante
            elif "FECHA" in celdas[0].text:
                celdas[1].text = fecha.strftime("%d.%m.%Y")
            elif "ALUMNO" in celdas[0].text:
                celdas[1].text = alumno
            elif "PUESTO" in celdas[0].text:
                celdas[1].text = puesto
            elif celdas[0].text.strip() in notas_categoria:
                letra = notas_categoria[celdas[0].text.strip()]
                idx = {"A": 1, "B": 2, "C": 3, "D": 4}[letra] - 1
                par = celdas[idx].paragraphs[0]
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                par.text = "X"
            elif "Nota media:" in celdas[0].text:
                celdas[0].text = f"Nota media: {nota_media}"
            elif "OBSERVACIONES GENERAL" in celdas[0].text.upper():
                celdas[0].text = "OBSERVACIONES GENERAL / JUSTIFICACIÓN"
                celdas[1].text = observaciones

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if st.button("📄 Generar Informe Word"):
    archivo = generar_word()
    st.download_button(
        label="📥 Descargar Informe",
        data=archivo,
        file_name=f"Informe_{alumno.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
