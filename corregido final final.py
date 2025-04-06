import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from io import BytesIO

st.set_page_config(page_title="Informe Guardia CAES", layout="centered")
st.title("Informe Personal de Guardia - CAES")

# ENCABEZADO
st.header("Encabezado del Informe")
informante = st.text_input("Informante")
fecha = st.date_input("Fecha", value=date.today())
alumno = st.text_input("Alumno evaluado")
puesto = st.text_input("Puesto durante la guardia")

# EVALUACIÓN
st.header("Evaluación por Categorías")

categorias = {
    "POLICÍA": ["Pregunta 1", "Pregunta 2"],
    "DISCIPLINA": ["Pregunta 1", "Pregunta 2"],
    "INTERES": ["Pregunta 1", "Pregunta 2"],
    "RESPONSABILIDAD": ["Pregunta 1", "Pregunta 2"],
    "INICIATIVA": ["Pregunta 1", "Pregunta 2"],
    "CONFIANZA EN SI MISMO": ["Pregunta 1", "Pregunta 2"],
    "ACTITUD CON LOS SUBORDINADOS": ["Pregunta 1", "Pregunta 2"],
    "ACTITUD CON EL MANDO": ["Pregunta 1", "Pregunta 2"],
    "COMPETENCIA / EFICACIA": ["Pregunta 1", "Pregunta 2"],
    "TRATO": ["Pregunta 1", "Pregunta 2"],
    "RESISTENCIA A LA FATIGA": ["Pregunta 1", "Pregunta 2"]
}

notas_categoria = {}
total_puntos = 0
total_preguntas = 0

for categoria, preguntas in categorias.items():
    st.subheader(categoria)
    respuestas = [st.checkbox(p, key=f"{categoria}_{i}") for i, p in enumerate(preguntas)]
    puntos = sum(respuestas)
    nota = round((puntos / len(preguntas)) * 10, 2)
    total_puntos += puntos
    total_preguntas += len(preguntas)

    if nota >= 9:
        letra = "A"
    elif nota >= 7:
        letra = "B"
    elif nota >= 5:
        letra = "C"
    else:
        letra = "D"

    notas_categoria[categoria] = letra
    st.write(f"Nota: {nota} — Calificación: {letra}")

# NOTA FINAL Y OBSERVACIONES
st.header("Resumen Final")
nota_media = round((total_puntos / total_preguntas) * 10, 2)
if nota_media >= 9:
    letra_final = "A"
elif nota_media >= 7:
    letra_final = "B"
elif nota_media >= 5:
    letra_final = "C"
else:
    letra_final = "D"

st.write(f"📊 Nota media: {nota_media}")
st.write(f"🔠 Calificación final: {letra_final}")
observaciones = st.text_area("Observaciones generales / Justificación")

# GENERADOR DE INFORME
def generar_word():
    plantilla = "INFORME EN BLANCO.docx"
    doc = Document(plantilla)

    for table in doc.tables:
        for row in table.rows:
            celdas = row.cells
            texto = celdas[0].text.strip()
            if texto == "INFORMANTE":
                celdas[1].text = ""
                celdas[1].paragraphs[0].add_run(informante)
            elif texto == "FECHA":
                celdas[1].text = ""
                celdas[1].paragraphs[0].add_run(fecha.strftime("%d.%m.%Y"))
            elif texto == "ALUMNO":
                celdas[1].text = ""
                celdas[1].paragraphs[0].add_run(alumno)
            elif texto == "PUESTO":
                celdas[1].text = ""
                celdas[1].paragraphs[0].add_run(puesto)
            elif texto in notas_categoria:
                letra = notas_categoria[texto]
                idx = {"A": 1, "B": 2, "C": 3, "D": 4}[letra]
                for i in range(1, 5):
                    celdas[i].text = ""
                par = celdas[idx].paragraphs[0]
                par.clear()
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                par.add_run("X")
            elif "Nota media:" in texto:
                celdas[0].text = f"Nota media: {nota_media}"
            elif texto == "":
                celdas[1].text = observaciones

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if st.button("📄 Generar Informe Word"):
    archivo = generar_word()
    st.download_button(
        label="📥 Descargar Informe",
        data=archivo,
        file_name=f"Informe_{alumno.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
